import os
import uuid
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.http import JsonResponse, HttpResponse
from rest_framework.decorators import api_view, parser_classes
from rest_framework.parsers import MultiPartParser
from PyPDF2 import PdfReader
from haystack import Document
from haystack.components.preprocessors import DocumentSplitter
from haystack.components.embedders import OpenAIDocumentEmbedder
from haystack.utils import Secret
from pinecone_store import document_store, get_document_store, reset_document_store
from .rfp_analyzer import RFPAnalyzer, analysis_cache
from asgiref.sync import async_to_sync
from .rfp_chatbot import RFPChatbot
from rest_framework.response import Response
from rest_framework import status
import json
import asyncio
from pinecone_store import Pinecone
import numpy as np
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.table import Table, TableStyleInfo
from datetime import datetime
from pinecone_store import reset_document_store
from pinecone_store import get_session_index_name, pc, index_name_base
from django.conf import settings
from django.core.cache import cache
import logging
import zipfile
import tempfile
import shutil
from openai import OpenAI
from pinecone_store import document_store as global_document_store

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Create a global analyzer instance using our Pinecone document store.
analyzer = RFPAnalyzer(vector_store=document_store)

def extract_text_from_file(file_path_or_file, file_type=None):
    """
    Extract text from a file (PDF, DOCX, XLSX, etc.).
    
    Args:
        file_path_or_file: Either a file path string or a file-like object
        file_type: Optional file type override (pdf, docx, xlsx)
        
    Returns:
        For PDFs: List of dictionaries with 'text' and 'page_number' keys
        For other files: String containing the extracted text
    """
    try:
        # Determine file type if not provided
        if not file_type:
            if isinstance(file_path_or_file, str):
                file_name = file_path_or_file.lower()
            else:
                # For file-like objects, try to get the name
                file_name = getattr(file_path_or_file, 'name', '').lower()
                
            if file_name.endswith('.pdf'):
                file_type = 'pdf'
            elif file_name.endswith('.docx'):
                file_type = 'docx'
            elif file_name.endswith('.xlsx') or file_name.endswith('.xls'):
                file_type = 'excel'
            else:
                raise ValueError(f"Unsupported file type: {file_name}")
        
        # Extract text based on file type
        if file_type == 'pdf':
            return extract_text_from_pdf(file_path_or_file)
        elif file_type == 'docx':
            return extract_text_from_docx(file_path_or_file)
        elif file_type == 'excel':
            return extract_text_from_excel(file_path_or_file)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
            
    except Exception as e:
        print(f"Text extraction failed: {e}")
        import traceback
        print(traceback.format_exc())
        raise

def extract_text_from_pdf(file_path_or_file):
    """Extract text from PDF file with page tracking."""
    file_handle = None
    try:
        from PyPDF2 import PdfReader
        
        # Check if the input is a string (file path) or file-like object
        if isinstance(file_path_or_file, str):
            # It's a file path
            if not os.path.exists(file_path_or_file):
                raise FileNotFoundError(f"File not found at {file_path_or_file}")
                
            # Open the file and create a PDF reader
            file_handle = open(file_path_or_file, "rb")
            reader = PdfReader(file_handle)
        else:
            # Assume it's a file-like object (BytesIO)
            # Make a copy of the file content to avoid issues with closed files
            from io import BytesIO
            if hasattr(file_path_or_file, 'read'):
                content = file_path_or_file.read()
                file_handle = BytesIO(content)
                reader = PdfReader(file_handle)
            else:
                # If it's already a BytesIO or similar, use it directly
                reader = PdfReader(file_path_or_file)
        
        # Extract text with page tracking
        page_texts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                page_texts.append({
                    "text": page_text,
                    "page_number": i + 1
                })
        
        if not page_texts:
            raise ValueError("No extractable text found in PDF")
            
        return page_texts
        
    except Exception as e:
        print(f"PDF extraction failed: {e}")
        raise
    finally:
        # Make sure we close the file handle if we opened one
        if file_handle and hasattr(file_handle, 'close'):
            file_handle.close()

def extract_text_from_docx(file_path_or_file):
    """Extract text from Word document."""
    try:
        import docx
        
        # Check if the input is a string (file path) or file-like object
        if isinstance(file_path_or_file, str):
            # It's a file path
            if not os.path.exists(file_path_or_file):
                raise FileNotFoundError(f"File not found at {file_path_or_file}")
                
            # Open the file
            doc = docx.Document(file_path_or_file)
        else:
            # For file-like objects
            doc = docx.Document(file_path_or_file)
        
        # Extract text from paragraphs
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        extracted_text = "\n".join(full_text)
        
        if not extracted_text.strip():
            raise ValueError("No extractable text found in Word document")
            
        return extracted_text
        
    except Exception as e:
        print(f"Word extraction failed: {e}")
        raise

def extract_text_from_excel(file_path_or_file):
    """Extract text from Excel file."""
    try:
        import pandas as pd
        
        # Read all sheets
        if isinstance(file_path_or_file, str):
            # It's a file path
            if not os.path.exists(file_path_or_file):
                raise FileNotFoundError(f"File not found at {file_path_or_file}")
                
            # Read Excel file with pandas
            excel_file = pd.ExcelFile(file_path_or_file)
        else:
            # For file-like objects
            excel_file = pd.ExcelFile(file_path_or_file)
        
        # Extract text from all sheets
        full_text = []
        
        for sheet_name in excel_file.sheet_names:
            df = excel_file.parse(sheet_name)
            
            # Add sheet name as a header
            full_text.append(f"Sheet: {sheet_name}")
            
            # Convert headers to text
            headers = df.columns.tolist()
            full_text.append(" | ".join([str(h) for h in headers]))
            
            # Convert each row to text
            for _, row in df.iterrows():
                row_text = " | ".join([str(cell) for cell in row.values])
                full_text.append(row_text)
            
            full_text.append("\n")  # Add separator between sheets
        
        extracted_text = "\n".join(full_text)

        
        if not extracted_text.strip():
            raise ValueError("No extractable text found in Excel file")
            
        return extracted_text
        
    except Exception as e:
        print(f"Excel extraction failed: {e}")
        raise

@api_view(["POST"])
@parser_classes([MultiPartParser])
def upload_pdf(request):
    """
    Upload an RFP PDF file, extract text, split it into chunks,
    compute 1536-d OpenAI embeddings, and index them into Pinecone.
    """
    try:
        file = request.FILES.get("file")
        if not file or not file.name.endswith(".pdf"):
            return JsonResponse({"error": "Invalid file"}, status=400)

        # Generate a session ID if not provided
        session_id = request.data.get('session_id')
        if not session_id:
            session_id = str(uuid.uuid4())
            print(f"Generated new session ID: {session_id}")
            
        # Get the document store for this session
        document_store = get_document_store(session_id)
        print(f"Using document store for session: {session_id}")

        # Generate a unique identifier for this document
        unique_id = str(uuid.uuid4())
        file_path = f"rfp_documents/{unique_id}_{file.name}"
        file_name = default_storage.save(file_path, ContentFile(file.read()))
        print(f"Saved PDF at: {default_storage.path(file_name)}")

        # Extract text from the PDF
        try:
            extracted_text = extract_text_from_pdf(default_storage.path(file_name))
        except Exception as e:
            return JsonResponse({"error": f"Failed to read PDF: {str(e)}"}, status=500)

        # Split the text into document chunks
        splitter = DocumentSplitter(split_by="sentence", split_length=3, split_overlap=1)
        splitter.warm_up()
        docs = [Document(content=extracted_text)]
        # Add metadata
        for doc in docs:
            doc.metadata = {"filename": file.name}
            
        split_docs = splitter.run(docs)["documents"]
        
        # Add metadata to split documents
        for doc in split_docs:
            if not hasattr(doc, 'metadata') or doc.metadata is None:
                doc.metadata = {}
            doc.metadata["filename"] = file.name
            
        print(f"Split into {len(split_docs)} document chunks")

        # Get OpenAI API key
        openai_api_key = os.environ.get("OPENAI_API_KEY")
        if not openai_api_key:
            return JsonResponse(
                {"error": "OPENAI_API_KEY environment variable is not set."},
                status=500,
            )

        # Embed documents
        document_embedder = OpenAIDocumentEmbedder(
            api_key=Secret.from_token(openai_api_key),
            model="text-embedding-ada-002"
        )
        embedding_results = document_embedder.run(split_docs)
        embedded_docs = embedding_results["documents"]

        # Debug: Log embedding dimensions
        for i, doc in enumerate(embedded_docs):
            if doc.embedding:
                print(f"Document {i} embedding dimension: {len(doc.embedding)}")
            else:
                print(f"Document {i} has no embedding.")

        # Write to Pinecone
        document_store.write_documents(embedded_docs)

        return JsonResponse({
            "success": True,
            "message": "Document uploaded and indexed successfully",
            "doc_id": unique_id,
            "session_id": session_id
        })

    except Exception as e:
        import traceback
        print(f"Error in upload_pdf: {str(e)}")
        print(traceback.format_exc())
        return JsonResponse({"error": str(e)}, status=500)

@api_view(["POST"])
@parser_classes([MultiPartParser])
def analyze_documents(request):
    """Process and index documents in Pinecone."""
    try:
        # Log request information
        logger.info(f"Request data keys: {request.data.keys()}")
        logger.info(f"Request FILES keys: {request.FILES.keys()}")
        
        # Generate a session ID if not provided
        session_id = request.data.get('session_id')
        if not session_id:
            session_id = str(uuid.uuid4())
            logger.info(f"Generated new session ID: {session_id}")
        
        # Clear any cached analysis for this session
        if session_id in analysis_cache:
            del analysis_cache[session_id]
        
        # Get the document store for this session
        document_store = get_document_store(session_id)
        logger.info(f"Using document store for session: {session_id}")
        
        # Check if file is in the request - try both 'file' and 'files' keys
        if 'file' in request.FILES:
            uploaded_file = request.FILES['file']
        elif 'files' in request.FILES:
            uploaded_file = request.FILES['files']
        else:
            # Try to find any file in the request
            all_files = list(request.FILES.values())
            if all_files:
                # Use the first file found
                uploaded_file = all_files[0]
                logger.info(f"File not found under 'file' key, but found: {uploaded_file.name}")
            else:
                logger.error("No files found in request")
                return JsonResponse({
                    "error": "No file provided",
                    "request_data_keys": list(request.data.keys()),
                    "request_files_keys": list(request.FILES.keys())
                }, status=400)
        
        logger.info(f"Processing file: {uploaded_file.name}, size: {uploaded_file.size}")
        
        # Initialize split_docs list
        split_docs = []
        
        # Process the file based on its type
        if uploaded_file.name.lower().endswith('.zip'):
            logger.info(f"Processing ZIP file: {uploaded_file.name}")
            
            # Create a temporary directory for extraction
            temp_dir = tempfile.mkdtemp()
            logger.info(f"Created temporary directory: {temp_dir}")
            
            try:
                # Save the zip file temporarily
                zip_path = os.path.join(temp_dir, uploaded_file.name)
                with open(zip_path, 'wb') as f:
                    for chunk in uploaded_file.chunks():
                        f.write(chunk)
                
                logger.info(f"Saved ZIP file to: {zip_path}")
                
                # Extract the zip file
                extracted_files = []
                
                try:
                    with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                        # Get list of all files in the zip
                        for file_info in zip_ref.infolist():
                            # Skip directories and macOS hidden files
                            if (file_info.is_dir() or 
                                file_info.filename.startswith('__MACOSX/') or
                                os.path.basename(file_info.filename).startswith('._') or
                                os.path.basename(file_info.filename) == '.DS_Store'):
                                logger.info(f"Skipping file/directory: {file_info.filename}")
                                continue
                                
                            # Extract the file
                            file_path = os.path.join(temp_dir, file_info.filename)
                            zip_ref.extract(file_info, temp_dir)
                            
                            # Add to our list
                            extracted_files.append({
                                'path': file_path,
                                'name': os.path.basename(file_info.filename)
                            })
                    
                    logger.info(f"Extracted {len(extracted_files)} files from ZIP")
                except Exception as zip_error:
                    logger.error(f"Error extracting ZIP file: {str(zip_error)}")
                    return JsonResponse({"error": f"Failed to extract ZIP file: {str(zip_error)}"}, status=400)
                
                # Process each file
                for file_data in extracted_files:
                    file_path = file_data['path']
                    file_name = file_data['name']
                    
                    # Skip macOS hidden files and other files we can't process
                    if (file_name.startswith('._') or 
                        file_name.startswith('.DS_Store') or
                        not (file_name.lower().endswith('.pdf') or 
                             file_name.lower().endswith('.docx') or 
                             file_name.lower().endswith('.xlsx') or
                             file_name.lower().endswith('.xls'))):
                        logger.info(f"Skipping file: {file_name}")
                        continue
                    
                    logger.info(f"Processing file from ZIP: {file_name}")
                    
                    try:
                        # Determine file type
                        file_type = None
                        if file_name.lower().endswith('.pdf'):
                            file_type = 'pdf'
                        elif file_name.lower().endswith('.docx'):
                            file_type = 'docx'
                        elif file_name.lower().endswith(('.xlsx', '.xls')):
                            file_type = 'excel'
                        
                        # Extract text from the file
                        extracted_text = extract_text_from_file(file_path, file_type)
                        logger.info(f"Extracted text from {file_name}")
                        
                        # Create a document splitter
                        splitter = DocumentSplitter(split_by="sentence", split_length=20, split_overlap=2)
                        splitter.warm_up()
                        
                        # Create documents based on file type
                        if file_type == 'pdf':
                            # For PDFs, create documents with page metadata
                            docs = []
                            for page_info in extracted_text:
                                doc = Document(content=page_info["text"])
                                doc.metadata = {
                                    "filename": file_name,
                                    "source": f"ZIP: {uploaded_file.name}",
                                    "page_number": page_info["page_number"]
                                }
                                docs.append(doc)
                        else:
                            # For other file types, create a single document
                            docs = [Document(content=extracted_text)]
                            # Add metadata
                            for doc in docs:
                                doc.metadata = {
                                    "filename": file_name,
                                    "source": f"ZIP: {uploaded_file.name}"
                                }
                        
                        # Split the documents
                        for doc in docs:
                            # Split each document and preserve its metadata
                            doc_splits = splitter.run([doc])["documents"]
                            
                            # Ensure metadata is preserved in each split
                            for split_doc in doc_splits:
                                if not hasattr(split_doc, 'metadata') or split_doc.metadata is None:
                                    split_doc.metadata = {}
                                
                                # Copy metadata from parent document
                                if hasattr(doc, 'metadata') and doc.metadata:
                                    for key, value in doc.metadata.items():
                                        split_doc.metadata[key] = value
                            
                            split_docs.extend(doc_splits)
                        
                        logger.info(f"Added {len(doc_splits)} chunks from {file_name}")
                        
                    except Exception as e:
                        logger.error(f"Error processing file {file_name}: {str(e)}")
                        # Continue with other files even if one fails
                
            finally:
                # Clean up temporary directory
                try:
                    shutil.rmtree(temp_dir)
                    logger.info("Cleaned up temporary directory")
                except Exception as cleanup_error:
                    logger.error(f"Error cleaning up temporary directory: {str(cleanup_error)}")
            
            # If no documents were processed successfully
            if not split_docs:
                logger.error("No valid documents found in ZIP file")
                return JsonResponse({"error": "No valid documents found in ZIP file"}, status=400)
            
        else:
            # Process a single file
            logger.info(f"Processing single file: {uploaded_file.name}")
            
            # Save the file temporarily
            file_path = default_storage.save(f"uploads/{uploaded_file.name}", ContentFile(uploaded_file.read()))
            logger.info(f"Saved file to: {file_path}")
            
            # Determine file type
            file_type = None
            if uploaded_file.name.lower().endswith('.pdf'):
                file_type = 'pdf'
            elif uploaded_file.name.lower().endswith('.docx'):
                file_type = 'docx'
            elif uploaded_file.name.lower().endswith(('.xlsx', '.xls')):
                file_type = 'excel'
            else:
                logger.error(f"Unsupported file type: {uploaded_file.name}")
                return JsonResponse({"error": "Unsupported file type"}, status=400)
            
            # Extract text from the file
            extracted_text = extract_text_from_file(default_storage.path(file_path), file_type)
            
            # Create a document splitter
            splitter = DocumentSplitter(split_by="sentence", split_length=20, split_overlap=2)
            # Warm up the splitter before using it
            splitter.warm_up()
            
            # Create documents based on file type
            if file_type == 'pdf':
                # For PDFs, create documents with page metadata
                docs = []
                for page_info in extracted_text:
                    doc = Document(content=page_info["text"])
                    doc.metadata = {
                        "filename": uploaded_file.name,
                        "page_number": page_info["page_number"]
                    }
                    docs.append(doc)
            else:
                # For other file types, create a single document
                docs = [Document(content=extracted_text)]
                # Add metadata
                for doc in docs:
                    doc.metadata = {"filename": uploaded_file.name}
            
            # Split the documents
            for doc in docs:
                # Split each document and preserve its metadata
                doc_splits = splitter.run([doc])["documents"]
                
                # Ensure metadata is preserved in each split
                for split_doc in doc_splits:
                    if not hasattr(split_doc, 'metadata') or split_doc.metadata is None:
                        split_doc.metadata = {}
                    
                    # Copy metadata from parent document
                    if hasattr(doc, 'metadata') and doc.metadata:
                        for key, value in doc.metadata.items():
                            split_doc.metadata[key] = value
                
                split_docs.extend(doc_splits)
            
            # Clean up the temporary file
            default_storage.delete(file_path)
        
        # Add page numbers to document content
        for doc in split_docs:
            if hasattr(doc, 'metadata') and doc.metadata and 'page_number' in doc.metadata:
                page_number = doc.metadata['page_number']
                # Add page number as a prefix to the content if not already present
                if not doc.content.startswith(f"[Page {page_number}]"):
                    doc.content = f"[Page {page_number}] {doc.content}"
                    logger.info(f"Added page number {page_number} to document content")

        # Log a sample document to verify
        if split_docs:
            logger.info(f"Sample document content with page number: {split_docs[0].content[:100]}...")
            logger.info(f"Sample document metadata: {split_docs[0].metadata}")
        
        # Get OpenAI API key - use only the dedicated key without fallback
        api_key = os.getenv("BID_QUALIFIER_OPENAI_API_KEY")
        if api_key:
            # Only log the first 5 and last 4 characters for security
            masked_key = api_key[:5] + "..." + api_key[-4:]
            logger.info(f"Using BID_QUALIFIER_OPENAI_API_KEY for document embedding: {masked_key}")
        
        if not api_key:
            logger.error("BID_QUALIFIER_OPENAI_API_KEY not found!")
            return JsonResponse({"error": "BID_QUALIFIER_OPENAI_API_KEY not found"}, status=500)
        
        # Embed the documents with the dedicated key
        logger.info(f"Embedding {len(split_docs)} document chunks")
        embedder = OpenAIDocumentEmbedder(api_key=Secret.from_token(api_key))
        embedded_docs = embedder.run(split_docs)["documents"]
        
        # Write to Pinecone
        document_store.write_documents(embedded_docs)
        logger.info(f"Indexed {len(embedded_docs)} document chunks in Pinecone")
        
        return JsonResponse({
            "success": True,
            "message": f"Documents analyzed and indexed successfully ({len(embedded_docs)} chunks)",
            "session_id": session_id
        })
        
    except Exception as e:
        import traceback
        logger.error(f"Error in analyze_documents: {str(e)}")
        logger.error(traceback.format_exc())
        return JsonResponse({
            "error": f"Analysis failed: {str(e)}"
        }, status=500)

@api_view(["POST"])
def analyze_rfp(request):
    """Analyze an RFP document."""
    try:
        # Get the session ID and template type from the request
        data = json.loads(request.body)
        session_id = data.get('session_id')
        template_type = data.get('template_type', 'standard')
        
        if not session_id:
            return JsonResponse({"error": "No session ID provided"}, status=400)
            
        # Get the document store for this session
        document_store = get_document_store(session_id)
        
        # Initialize the analyzer with the document store
        analyzer = RFPAnalyzer(vector_store=document_store)
        
        # Analyze the RFP with the specified template
        analysis = async_to_sync(analyzer.analyze_rfp)(
            text="Extract all key information from this RFP document.",
            template_type=template_type
        )
        
        return JsonResponse({
            "success": True,
            "result": analysis,
            "session_id": session_id,
            "template_used": template_type
        })
        
    except Exception as e:
        import traceback
        print(f"Error in analyze_rfp: {str(e)}")
        print(traceback.format_exc())
        return JsonResponse({
            "error": f"Analysis failed: {str(e)}"
        }, status=500)

@api_view(["POST"])
def generate_bid_matrix(request, doc_id):
    """
    Generate a bid matrix based on an RFP document.
    """
    result = async_to_sync(analyzer.generate_bid_matrix)({"doc_id": doc_id})
    return JsonResponse({"matrix": result})

@api_view(["GET"])
def download_matrix(request, doc_id):
    """
    Download the bid matrix as an Excel file.
    """
    import io
    import pandas as pd 

    data = {
        "Requirement": [],
        "Priority": [],
        "Complexity": [],
        "Status": [],
        "Assigned To": [],
        "Notes": [],
    }
    df = pd.DataFrame(data)
    excel_buffer = io.BytesIO()
    df.to_excel(excel_buffer, index=False)
    excel_buffer.seek(0)

    response = HttpResponse(
        excel_buffer.read(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = f"attachment; filename=bid_matrix_{doc_id}.xlsx"
    return response

@api_view(['POST'])
def chat_with_rfp(request):
    """
    Chat with the RFP using the session data.
    """
    try:
        # Log the request data for debugging
        print("Request data:", request.data)
        
        # Get the session ID from the request
        session_id = request.data.get('session_id')
        message = request.data.get('message')
        
        print("Extracted session_id:", session_id)
        print("Extracted message:", message)
        
        if not message:
            return JsonResponse({'error': 'Message is required'}, status=400)
        
        # Get the document store - use global store if no session ID
        if session_id:
            document_store = get_document_store(session_id)
        else:
            # Use the global document store if no session ID is provided
            document_store = global_document_store
        
        # Create a chatbot instance
        chatbot = RFPChatbot(document_store)
        
        # Get the response using the existing get_response method
        response = chatbot.get_response(message)
        
        # Extract the answer from the response
        if response and 'answer' in response:
            return JsonResponse({'response': response['answer']})
        else:
            return JsonResponse({'error': 'Failed to get a response from the chatbot'}, status=500)
    
    except Exception as e:
        logger.error(f"Error in chat_with_rfp: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

@api_view(["GET"])
def compare_indexes(request):
    """Compare documents in paidmediabids against session-specific RFP index for similarity."""
    try:
        # Get session ID from query parameters
        session_id = request.query_params.get('session_id')
        
        if not session_id:
            return JsonResponse({
                'success': False,
                'error': 'No session_id provided. Please include a session_id query parameter.'
            }, status=400)
        
        # Initialize Pinecone
        pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
        
        # Get the index name for this session
        rfp_index_name = get_session_index_name(session_id)
        print(f"Using RFP index: {rfp_index_name}")
        
        # Check if the index exists
        available_indexes = pc.list_indexes().names()
        if rfp_index_name not in available_indexes:
            return JsonResponse({
                'success': False,
                'error': f'Index "{rfp_index_name}" not found. Available indexes: {available_indexes}'
            }, status=404)
        
        # Get both indexes
        uswebbid = pc.Index("paidmediabids")
        rfp = pc.Index(rfp_index_name)
        
        # Get vectors from rfp-index
        rfp_response = rfp.query(
            vector=[1.0] * 1536,
            top_k=50,
            include_values=True,
            include_metadata=True,
            namespace="default"
        ).to_dict()

        # Get the first vector from rfp-index results
        if rfp_response.get('matches'):
            rfp_vector = rfp_response['matches'][0]['values']
            
            # Use this vector to query uswebbid
            uswebbid_response = uswebbid.query(
                vector=rfp_vector,
                top_k=50,
                include_values=True,
                include_metadata=True,
                namespace="default"
            ).to_dict()
            
            # Calculate average similarity score
            scores = [match['score'] for match in uswebbid_response.get('matches', [])]
            average_similarity = sum(scores) / len(scores) if scores else 0
            
            return JsonResponse({
                'success': True,
                'similarity_score': round(average_similarity, 4),
                'total_documents_compared': len(scores)
            })
        else:
            return JsonResponse({
                'success': False,
                'error': f'No vectors found in {rfp_index_name} index'
            })

    except Exception as e:
        print(f"Error: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)

@api_view(["POST"])
def download_report(request):
    """Download the RFP analysis as an Excel report."""
    try:
        print("Starting report generation...")
        
        # Get the RFP data from the request
        rfp_data = request.data.get('rfpData', {})
        print(f"Received RFP data keys: {rfp_data.keys()}")  # Debug log

        # Create a new workbook and worksheet
        wb = Workbook()
        ws = wb.active
        ws.title = "RFP Analysis Report"

        # Add title and date
        ws['A1'] = "RFP Analysis Report"
        ws['A2'] = f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        
        # Apply formatting to title
        title_font = Font(name='Arial', size=16, bold=True)
        ws['A1'].font = title_font
        
        # Apply formatting to date
        date_font = Font(name='Arial', size=12, italic=True)
        ws['A2'].font = date_font
        
        ws.merge_cells('A1:F1')
        ws.merge_cells('A2:F2')
        
        # Add header row
        headers = ["Section", "Field", "Value", "Confidence", "Interpreted", "Source Page"]
        for col, header in enumerate(headers, start=1):
            cell = ws.cell(row=4, column=col, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="E0E0E0", end_color="E0E0E0", fill_type="solid")
        
        # Set column widths
        ws.column_dimensions['A'].width = 20  # Section
        ws.column_dimensions['B'].width = 25  # Field
        ws.column_dimensions['C'].width = 40  # Value
        ws.column_dimensions['D'].width = 12  # Confidence
        ws.column_dimensions['E'].width = 12  # Interpreted
        ws.column_dimensions['F'].width = 15  # Source Page
        
        # Start from row 5 (after headers)
        current_row = 5
        
        # Process each section in the RFP data
        for section_name, section_data in rfp_data.items():
            if not isinstance(section_data, dict):
                continue
                
            # Format section name for display
            display_section = section_name.replace('_', ' ').title()
            
            # Process each field in the section
            for field_name, field_data in section_data.items():
                # Format field name for display
                display_field = field_name.replace('_', ' ').title()
                
                # Handle different data structures
                if isinstance(field_data, dict) and 'value' in field_data:
                    # New format with confidence and other metadata
                    value = field_data.get('value', '')
                    confidence = field_data.get('confidence', 0)
                    is_interpreted = field_data.get('is_interpreted', False)
                    source_page = field_data.get('source_page', '')
                    
                    # Write to Excel
                    ws.cell(row=current_row, column=1, value=display_section)
                    ws.cell(row=current_row, column=2, value=display_field)
                    ws.cell(row=current_row, column=3, value=value)
                    
                    # Format confidence as percentage
                    confidence_cell = ws.cell(row=current_row, column=4, value=confidence)
                    confidence_cell.number_format = '0.0%'
                    
                    # Color code confidence
                    if confidence >= 0.8:
                        confidence_cell.fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")  # Green
                    elif confidence >= 0.5:
                        confidence_cell.fill = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")  # Yellow
                    elif confidence > 0:
                        confidence_cell.fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")  # Red
                    
                    # Format is_interpreted as Yes/No
                    ws.cell(row=current_row, column=5, value="Yes" if is_interpreted else "No")
                    
                    # Source page
                    ws.cell(row=current_row, column=6, value=source_page)
                    
                else:
                    # Old format or simple value
                    ws.cell(row=current_row, column=1, value=display_section)
                    ws.cell(row=current_row, column=2, value=display_field)
                    ws.cell(row=current_row, column=3, value=str(field_data))
                    # Leave other columns empty
                
                current_row += 1
            
            # Add a blank row between sections
            current_row += 1
        
        # Create a table for easy filtering and sorting
        tab = Table(displayName="RFPAnalysisTable", ref=f"A4:{get_column_letter(len(headers))}{current_row-1}")
        style = TableStyleInfo(
            name="TableStyleMedium9", 
            showFirstColumn=False,
            showLastColumn=False, 
            showRowStripes=True, 
            showColumnStripes=False
        )
        tab.tableStyleInfo = style
        ws.add_table(tab)
        
        # Freeze the header row
        ws.freeze_panes = 'A5'
        
        # Create response
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = 'attachment; filename=rfp_analysis_report.xlsx'

        # Save workbook to response
        wb.save(response)
        print("Report generated successfully")
        return response

    except Exception as e:
        print(f"Error generating report: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)

@api_view(["POST"])
def cleanup_session(request):
    """Clean up a session's resources."""
    try:
        session_id = request.data.get('session_id')
        if not session_id:
            return JsonResponse({"error": "No session ID provided"}, status=400)
        
        # Clear any cached analysis for this session
        if session_id in analysis_cache:
            del analysis_cache[session_id]
        
        # Get the index name for this session
        index_name = get_session_index_name(session_id)
        print(f"Cleaning up session {session_id}, index name: {index_name}")
        
        # SAFETY CHECK: Only delete if it's a session index
        # This prevents deletion of static indexes
        if index_name.startswith(f"{index_name_base}-") and index_name != index_name_base:
            print(f"Index {index_name} is a session index, checking if it exists...")
            
            # List all indexes to check if this one exists
            existing_indexes = pc.list_indexes().names()
            print(f"Existing indexes: {existing_indexes}")
            
            if index_name in existing_indexes:
                # Check if the index is protected
                protected_indexes = getattr(settings, 'PROTECTED_INDEXES', [])
                print(f"Protected indexes: {protected_indexes}")
                
                if index_name in protected_indexes:
                    print(f"Index {index_name} is protected, not deleting")
                    return JsonResponse({
                        "error": "Cannot delete protected index",
                        "message": f"Index {index_name} is protected"
                    }, status=403)
                else:
                    print(f"Deleting index {index_name}")
                    pc.delete_index(index_name)
                    return JsonResponse({
                        "success": True,
                        "message": f"Session {session_id} cleaned up successfully"
                    })
            else:
                print(f"Index {index_name} not found in existing indexes")
                return JsonResponse({
                    "success": True,
                    "message": f"No index found for session {session_id}"
                })
        else:
            print(f"Index {index_name} is not a session index, not deleting")
            return JsonResponse({
                "error": "Cannot delete non-session index",
                "message": f"Index {index_name} appears to be a static index"
            }, status=403)
            
    except Exception as e:
        import traceback
        print(f"Error in cleanup_session: {str(e)}")
        print(traceback.format_exc())
        return JsonResponse({
            "error": f"Failed to clean up session: {str(e)}"
        }, status=500)

@api_view(["POST"])
@parser_classes([MultiPartParser])
def analyze_pdf(request):
    """Process and index a PDF in Pinecone (legacy endpoint, use analyze_documents instead)."""
    try:
        print("analyze_pdf view called (legacy endpoint)")
        print("Request data:", request.data)
        print("Request FILES:", request.FILES)
        
        # Generate a session ID if not provided
        session_id = request.data.get('session_id')
        if not session_id:
            session_id = str(uuid.uuid4())
            print(f"Generated new session ID: {session_id}")
        
        # Clear any cached analysis for this session
        if session_id in analysis_cache:
            del analysis_cache[session_id]
        
        # Get the document store for this session
        document_store = get_document_store(session_id)
        print(f"Using document store for session: {session_id}")
        
        # Get the file from the request
        if 'file' not in request.FILES:
            print("No file in request.FILES")
            return JsonResponse({"error": "No file provided"}, status=400)
        
        uploaded_file = request.FILES['file']
        print(f"Received file: {uploaded_file.name}, size: {uploaded_file.size}")
        
        # Check if it's a PDF
        if not uploaded_file.name.lower().endswith('.pdf'):
            return JsonResponse({"error": "File must be a PDF"}, status=400)
        
        # Save the file temporarily
        file_path = default_storage.save(f"uploads/{uploaded_file.name}", ContentFile(uploaded_file.read()))
        print(f"Saved file to: {file_path}")
        
        # Instead of getting an absolute path, read the file directly from storage
        with default_storage.open(file_path, 'rb') as f:
            pdf_content = f.read()
            
        # Use BytesIO instead of file path for PDF extraction
        from io import BytesIO
        pdf_file = BytesIO(pdf_content)
        
        # Extract and process the PDF using the BytesIO object
        extracted_text = extract_text_from_pdf(pdf_file)
        print(f"Extracted text length: {len(extracted_text)}")
        
        # Split into chunks and embed
        splitter = DocumentSplitter(split_by="sentence", split_length=3, split_overlap=1)
        # Warm up the splitter before using it
        splitter.warm_up()
        
        docs = [Document(content=extracted_text)]
        # Add metadata
        for doc in docs:
            doc.metadata = {"filename": uploaded_file.name}
            
        split_docs = splitter.run(docs)["documents"]
        
        # Add metadata to split documents
        for doc in split_docs:
            if not hasattr(doc, 'metadata') or doc.metadata is None:
                doc.metadata = {}
            doc.metadata["filename"] = uploaded_file.name

        # Add page numbers to document content
        for doc in split_docs:
            if hasattr(doc, 'metadata') and doc.metadata and 'page_number' in doc.metadata:
                page_number = doc.metadata['page_number']
                # Add page number as a prefix to the content if not already present
                if not doc.content.startswith(f"[Page {page_number}]"):
                    doc.content = f"[Page {page_number}] {doc.content}"
                    print(f"Added page number {page_number} to document content")

        # Log a sample document to verify
        if split_docs:
            print(f"Sample document content with page number: {split_docs[0].content[:100]}...")
            print(f"Sample document metadata: {split_docs[0].metadata}")

        # Get OpenAI API key - use only the dedicated key without fallback
        api_key = os.getenv("BID_QUALIFIER_OPENAI_API_KEY")
        if not api_key:
            return JsonResponse({"error": "BID_QUALIFIER_OPENAI_API_KEY not found"}, status=500)

        # Embed the documents with the dedicated key
        embedder = OpenAIDocumentEmbedder(api_key=Secret.from_token(api_key))
        embedded_docs = embedder.run(split_docs)["documents"]

        # Write to Pinecone
        document_store.write_documents(embedded_docs)

        # Clean up the temporary file
        default_storage.delete(file_path)

        return JsonResponse({
            "success": True,
            "message": "Document analyzed and indexed successfully",
            "session_id": session_id
        })

    except Exception as e:
        import traceback
        print(f"Error in analyze_pdf: {str(e)}")
        print(traceback.format_exc())
        return JsonResponse({
            "error": f"Analysis failed: {str(e)}"
        }, status=500)

@api_view(["POST"])
def clear_session(request):
    """Clear all documents from a session's index."""
    try:
        # Get the session ID
        session_id = request.data.get('session_id')
        if not session_id:
            return JsonResponse({"error": "No session ID provided"}, status=400)
        
        # Clear any cached analysis for this session
        if session_id in analysis_cache:
            del analysis_cache[session_id]
        
        # Reset the document store for this session
        document_store = reset_document_store(session_id)
        print(f"Reset document store for session: {session_id}")
        
        return JsonResponse({
            "success": True,
            "message": f"Session {session_id} cleared successfully",
            "session_id": session_id
        })

    except Exception as e:
        import traceback
        print(f"Error in clear_session: {str(e)}")
        print(traceback.format_exc())
        return JsonResponse({
            "error": f"Failed to clear session: {str(e)}"
        }, status=500)

@api_view(["GET"])
def check_model_limits(request):
    """Check and return the token limits for the GPT-4o model."""
    try:
        client = OpenAI(api_key=os.getenv("BID_QUALIFIER_OPENAI_API_KEY"))
        models = client.models.list()
        
        model_info = {}
        # Find GPT-4o in the list
        for model in models.data:
            if "gpt-4" in model.id:  # Get all GPT-4 variants
                model_info[model.id] = {
                    "context_window": getattr(model, 'context_window', 'Not specified'),
                    "max_tokens": getattr(model, 'max_tokens', 'Not specified')
                }
                
        return JsonResponse({
            "success": True,
            "models": model_info
        })
    except Exception as e:
        import traceback
        print(f"Error checking model limits: {e}")
        print(traceback.format_exc())
        return JsonResponse({
            "error": f"Failed to check model limits: {str(e)}"
        }, status=500)